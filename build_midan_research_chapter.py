from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\أشيائي\Correct Project Structure intel system")
OUT_DIR = ROOT / "research_chapter_output"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "MIDAN_Research_Chapter_Revised.docx"


COLORS = {
    "ink": (30, 39, 46),
    "muted": (85, 98, 112),
    "line": (114, 137, 150),
    "paper": (250, 252, 253),
    "blue": (35, 86, 135),
    "blue_soft": (223, 235, 246),
    "green": (44, 121, 82),
    "green_soft": (224, 241, 232),
    "amber": (158, 111, 35),
    "amber_soft": (247, 237, 218),
    "red": (159, 70, 70),
    "red_soft": (247, 226, 226),
    "gray_soft": (235, 239, 242),
}


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        trial = word if not line else f"{line} {word}"
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: tuple[int, int, int] = COLORS["blue_soft"],
    outline: tuple[int, int, int] = COLORS["blue"],
    title_color: tuple[int, int, int] = COLORS["ink"],
):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    tf = font(26, True)
    bf = font(20)
    draw.text((x1 + 22, y1 + 18), title, font=tf, fill=title_color)
    if body:
        y = y1 + 56
        for line in wrap(draw, body, x2 - x1 - 44, bf):
            draw.text((x1 + 22, y), line, font=bf, fill=COLORS["muted"])
            y += 25


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=COLORS["line"], width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - sign * 18, y2 - 9), (x2 - sign * 18, y2 + 9)]
    else:
        sign = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - sign * 18), (x2 + 9, y2 - sign * 18)]
    draw.polygon(pts, fill=color)


def title(draw: ImageDraw.ImageDraw, text: str, subtitle: str | None = None):
    draw.text((60, 35), text, font=font(34, True), fill=COLORS["ink"])
    if subtitle:
        draw.text((60, 78), subtitle, font=font(21), fill=COLORS["muted"])


def save_architecture():
    img = Image.new("RGB", (2200, 1500), COLORS["paper"])
    d = ImageDraw.Draw(img)
    title(d, "MIDAN System Design", "End-to-end agentic pipeline and evidence flow")
    boxes = [
        ((80, 160, 510, 285), "Agent 1", "Intake and validity gate: rejects impossible, broken, vague, or adversarial inputs.", COLORS["red_soft"], COLORS["red"]),
        ((640, 160, 1070, 285), "Agent 2", "Feature extraction: sector, country, model, segment, stage, readiness, risk, differentiation.", COLORS["blue_soft"], COLORS["blue"]),
        ((1200, 160, 1630, 285), "Agent 3", "Market intelligence: macro vector, SVM regime, FCM membership, SHAP, SARIMA.", COLORS["green_soft"], COLORS["green"]),
        ((80, 430, 510, 555), "Agent 4", "Structured reasoning: differentiation, competition, business model, unit economics, signal interactions.", COLORS["blue_soft"], COLORS["blue"]),
        ((640, 430, 1070, 555), "Agent 5", "Mechanism extraction: structural observations, moats, constraints, tensions, uncertainty.", COLORS["amber_soft"], COLORS["amber"]),
        ((1200, 430, 1630, 555), "Agent 6", "Precedent routing: attribution consistency, vector retrieval, novelty gate, route selection.", COLORS["green_soft"], COLORS["green"]),
        ((80, 700, 510, 825), "Agent 7", "Decision engine: risk decomposition, conflict detection, offsetting, decision state.", COLORS["red_soft"], COLORS["red"]),
        ((640, 700, 1070, 825), "Agent 8", "Synthesis and interaction: anchored explanation, response schema, chat behavior.", COLORS["blue_soft"], COLORS["blue"]),
        ((1200, 700, 1630, 825), "Agent 9", "Monitoring and feedback: prediction log, drift checks, outcome calibration.", COLORS["gray_soft"], COLORS["line"]),
    ]
    for b in boxes:
        box(d, *b)
    arrow(d, (510, 222), (640, 222))
    arrow(d, (1070, 222), (1200, 222))
    arrow(d, (1415, 285), (1415, 430))
    arrow(d, (1200, 492), (1070, 492))
    arrow(d, (640, 492), (510, 492))
    arrow(d, (295, 555), (295, 700))
    arrow(d, (510, 762), (640, 762))
    arrow(d, (1070, 762), (1200, 762))
    arrow(d, (1415, 825), (1415, 1030))
    arrow(d, (1415, 1030), (295, 1030))
    arrow(d, (295, 1030), (295, 825))
    d.rounded_rectangle((1760, 160, 2100, 825), radius=18, fill=(255, 255, 255), outline=COLORS["line"], width=3)
    d.text((1790, 190), "Primary Artifacts", font=font(26, True), fill=COLORS["ink"])
    artifact_lines = [
        "Structured response payload",
        "Risk decomposition",
        "Decision state",
        "Strategic anchors",
        "Mechanism envelope",
        "Prediction log",
        "Outcome log",
    ]
    y = 245
    for line in artifact_lines:
        d.ellipse((1790, y + 7, 1805, y + 22), fill=COLORS["blue"])
        d.text((1820, y), line, font=font(21), fill=COLORS["muted"])
        y += 55
    d.text((70, 1375), "Figure 1. Hand-made architecture diagram derived from the MIDAN implementation.", font=font(20), fill=COLORS["muted"])
    out = FIG_DIR / "figure_1_architecture.png"
    img.save(out, quality=95)
    return out


def save_market_flow():
    img = Image.new("RGB", (1900, 980), COLORS["paper"])
    d = ImageDraw.Draw(img)
    title(d, "Agent 3 Market Intelligence Flow", "Macro classification and explainability path")
    nodes = [
        ((80, 220, 380, 350), "Macro Lookup", "Country and sector tables"),
        ((500, 220, 800, 350), "Idea Deltas", "Confidence-gated adjustments"),
        ((920, 220, 1220, 350), "SVM + Rules", "Regime classification"),
        ((1340, 220, 1640, 350), "FCM", "Fuzzy membership"),
        ((920, 520, 1220, 650), "SHAP", "Top macro drivers"),
        ((1340, 520, 1640, 650), "SARIMA", "Sector trend signal"),
    ]
    for xy, t, b in nodes:
        box(d, xy, t, b, COLORS["green_soft"], COLORS["green"])
    arrow(d, (380, 285), (500, 285))
    arrow(d, (800, 285), (920, 285))
    arrow(d, (1220, 285), (1340, 285))
    arrow(d, (1070, 350), (1070, 520))
    arrow(d, (1220, 585), (1340, 585))
    box(d, (610, 760, 1510, 895), "Routing Composite", "Intelligent Score combines regime favorability, SVM margin, FCM membership, SARIMA trend, and SHAP attribution consistency.", COLORS["amber_soft"], COLORS["amber"])
    arrow(d, (1490, 350), (1110, 760))
    arrow(d, (1490, 650), (1180, 760))
    arrow(d, (1070, 650), (1070, 760))
    out = FIG_DIR / "figure_2_market_flow.png"
    img.save(out, quality=95)
    return out


def save_mechanism_flow():
    img = Image.new("RGB", (1900, 1150), COLORS["paper"])
    d = ImageDraw.Draw(img)
    title(d, "Agent 5 Mechanism Extraction", "From structural observations to uncertainty-aware mechanism evidence")
    labels = [
        ("Extractability", "Can mechanism inference run?"),
        ("Observations", "Translate idea and reasoning fields"),
        ("Assignment", "Map signals to mechanism types"),
        ("Calibration", "Set confidence and evidence strength"),
        ("Market Structure", "Classify structural context"),
        ("Tensions", "Detect mechanism conflicts"),
        ("Replication", "Estimate difficulty to copy"),
        ("Uncertainty", "Propagate evidence limits to Agent 7"),
    ]
    coords = [(80, 190), (500, 190), (920, 190), (1340, 190), (80, 570), (500, 570), (920, 570), (1340, 570)]
    for (t, b), (x, y) in zip(labels, coords):
        box(d, (x, y, x + 300, y + 135), t, b, COLORS["amber_soft"], COLORS["amber"])
    for i in range(3):
        arrow(d, (coords[i][0] + 300, coords[i][1] + 68), (coords[i + 1][0], coords[i + 1][1] + 68))
    arrow(d, (1490, 325), (230, 570))
    for i in range(4, 7):
        arrow(d, (coords[i][0] + 300, coords[i][1] + 68), (coords[i + 1][0], coords[i + 1][1] + 68))
    box(d, (500, 920, 1220, 1045), "Mechanism Envelope", "Mechanisms, market structure, tensions, replication profiles, uncertainty, consistency report, and epistemic summary.", COLORS["blue_soft"], COLORS["blue"])
    arrow(d, (1490, 705), (1220, 982))
    out = FIG_DIR / "figure_3_mechanism_flow.png"
    img.save(out, quality=95)
    return out


def save_decision_flow():
    img = Image.new("RGB", (1900, 1080), COLORS["paper"])
    d = ImageDraw.Draw(img)
    title(d, "Agent 7 Decision State Machine", "How MIDAN converts evidence into bounded decisions")
    top = [
        ((90, 190, 450, 330), "Risk Decomposition", "Market, execution, and timing risks are assessed separately."),
        ((580, 190, 940, 330), "Conflict Detection", "Contradictory signals are identified and severity-ranked."),
        ((1070, 190, 1430, 330), "Offsetting", "Strong evidence can reduce, but not erase, risk."),
    ]
    for xy, t, b in top:
        box(d, xy, t, b, COLORS["red_soft"], COLORS["red"])
    arrow(d, (450, 260), (580, 260))
    arrow(d, (940, 260), (1070, 260))
    box(d, (700, 465, 1080, 610), "Decision Quality", "Input completeness, signal agreement, assumption density, uncertainty.", COLORS["amber_soft"], COLORS["amber"])
    arrow(d, (1250, 330), (890, 465))
    states = [
        ((90, 790, 320, 910), "GO", "Low blocking risk"),
        ((390, 790, 620, 910), "CONDITIONAL", "Specific validation needed"),
        ((690, 790, 920, 910), "NO_GO", "Risk remains dominant"),
        ((990, 790, 1220, 910), "HIGH_UNCERTAINTY", "Advisory only"),
        ((1290, 790, 1560, 910), "CONFLICTING_SIGNALS", "Unresolved high conflict"),
    ]
    for xy, t, b in states:
        box(d, xy, t, b, COLORS["gray_soft"], COLORS["line"])
        arrow(d, (890, 610), ((xy[0] + xy[2]) // 2, xy[1]))
    out = FIG_DIR / "figure_4_decision_state.png"
    img.save(out, quality=95)
    return out


def save_data_pipeline():
    img = Image.new("RGB", (1900, 980), COLORS["paper"])
    d = ImageDraw.Draw(img)
    title(d, "Supporting Data Acquisition Pipeline", "Collection, extraction, validation, and structured startup intelligence")
    labels = [
        ("Sources", "Websites, YC, Failory, Reddit, Product Hunt"),
        ("Collectors", "Source-specific raw capture"),
        ("Extractors", "Normalize startup signals"),
        ("Validation", "Schema checks and cleaning"),
        ("Conflict Resolution", "Merge duplicate startup records"),
        ("Pattern Analysis", "Behavioral signal summaries"),
        ("Structured Corpus", "Startup intelligence JSON"),
    ]
    x = 80
    y = 260
    last = None
    for idx, (t, b) in enumerate(labels):
        xy = (x, y, x + 230, y + 150)
        box(d, xy, t, b, COLORS["blue_soft"] if idx < 3 else COLORS["green_soft"], COLORS["blue"] if idx < 3 else COLORS["green"])
        if last:
            arrow(d, (last[2], (last[1] + last[3]) // 2), (xy[0], (xy[1] + xy[3]) // 2))
        last = xy
        x += 260
    out = FIG_DIR / "figure_5_data_pipeline.png"
    img.save(out, quality=95)
    return out


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    for chunk in text.split("**"):
        r = p.add_run(chunk)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(30, 39, 46)
    return p


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for name, size in [("Heading 1", 15.5), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True


def build_doc():
    OUT_DIR.mkdir(exist_ok=True)
    FIG_DIR.mkdir(exist_ok=True)
    figs = [
        save_architecture(),
        save_market_flow(),
        save_mechanism_flow(),
        save_decision_flow(),
        save_data_pipeline(),
    ]

    doc = Document()
    style_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MIDAN: A Multi-Agent Market Intelligence and Decision Analysis Network for Startup Evaluation")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Publication-ready research chapter generated from the inspected project implementation")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    doc.add_paragraph()

    add_heading(doc, "1. Abstract", 1)
    add_para(doc, (
        "Startup evaluation is often performed with incomplete information, inconsistent market evidence, and excessive reliance on broad scoring heuristics. "
        "This chapter presents MIDAN, a multi-agent market intelligence and decision analysis network designed to evaluate early-stage startup ideas through explicit evidence gates, macroeconomic signal modeling, structured reasoning, and bounded decision synthesis. "
        "The system combines confidence-aware feature extraction, a market intelligence agent using regime classification and explainability signals, a structured reasoning agent for business-model and execution analysis, a mechanism extraction agent for competitive structure, a precedent-routing agent, and a final decision agent that separates market, execution, and timing risks. "
        "Rather than allowing a single numerical score to determine the verdict, MIDAN uses scores as routing evidence while preserving a discrete decision state as the authoritative output. "
        "Inspection of the implemented system shows that the architecture supports fail-fast clarification, transparent macro adjustments, traceable reasoning anchors, uncertainty-aware mechanism extraction, and strict response schema enforcement. "
        "Validation of the Phase 5 components produced 39 passing checks across intelligent-score computation, routing, drift monitoring, and decision integration. "
        "Local runtime scenarios further show differentiated behavior across viable, uncertain, and under-defined inputs. "
        "The contribution is a grounded system architecture for startup evaluation that combines machine learning signals with explicit epistemic controls and human-readable decision logic."
    ))

    add_heading(doc, "2. Introduction", 1)
    intro = [
        "Early-stage startup evaluation sits between two kinds of uncertainty. The first is market uncertainty: the economic environment, sector trend, competitive density, and available capital may support or suppress a new venture independently of the idea itself. The second is concept uncertainty: the idea may lack a clear user, a specific pain point, a viable business model, or a defensible mechanism. Most lightweight evaluation tools collapse these uncertainties into a score. That approach is convenient, but it hides the reason for the recommendation and can make a weak signal appear more precise than it is.",
        "MIDAN was built to address that problem as an implemented decision system. Its central premise is simple: a startup idea should not receive a strategic verdict until the system can identify what the idea is, what market context it belongs to, which assumptions drive its viability, and where the evidence is missing. This is why the system uses gates before prediction, separates macro evidence from idea-level reasoning, and treats the final decision as a state-machine output rather than a raw numerical score.",
        "The project is implemented as a Python package with a FastAPI interface, a Streamlit dashboard, a static frontend, trained model artifacts, structured response schemas, runtime logging, and a supporting data acquisition subsystem. The runtime system is organized around agents, each with a bounded responsibility. The agents do not simply pass text to a language model. They perform validity checking, confidence-scored parsing, macro vector construction, regime classification, fuzzy membership analysis, explainability scoring, time-series trend lookup, mechanism extraction, precedent routing, decision-state derivation, and calibrated synthesis.",
        "The design is intentionally conservative. MIDAN blocks vague or contradictory inputs before inference. It marks unknown fields explicitly. It applies idea-derived macro adjustments only when source fields exceed a confidence floor. It treats retrieval-based precedent as advisory. It forces human review when routing signals conflict too strongly. It logs prediction signals for drift monitoring and stores outcomes separately for later calibration. These choices make the system less theatrical but more defensible.",
        "This chapter revises the system description into a research-style account of the actual implementation. It avoids treating MIDAN as a product pitch or a documentation dump. The emphasis is on architecture, information flow, evidence boundaries, implemented behavior, and validation signals available in the project. The chapter also states limitations directly, including current data staleness, artifact availability, narrow logged runtime coverage, and the need for broader empirical validation before deployment in high-stakes investment workflows.",
    ]
    for para in intro:
        add_para(doc, para)

    doc.add_page_break()
    add_heading(doc, "3. System Design", 1)
    doc.add_picture(str(figs[0]), width=Inches(7.2))
    add_caption(doc, "Figure 1. System-level MIDAN architecture. The diagram is hand-made from the inspected implementation and does not reuse model-generated figures.")

    agent_rows = [
        ["Agent 1", "Intake and Validity Agent", "Screens the raw idea before inference. It blocks impossible, contradictory, spam-like, prompt-injected, and under-defined inputs so later agents do not reason over unusable evidence."],
        ["Agent 2", "Feature Extraction Agent", "Converts text into a confidence-scored idea schema. Required fields include business model, target segment, and stage, with low-confidence values surfaced as unknown rather than silently guessed."],
        ["Agent 3", "Market Intelligence Agent", "Constructs a macro vector from country and sector signals, applies confidence-gated idea adjustments, classifies the market regime, and produces explainability and trend signals."],
        ["Agent 4", "Structured Reasoning Agent", "Interprets the idea through differentiation, competition, business model, unit economics, and signal interaction modules. It exposes insufficient information at the module level."],
        ["Agent 5", "Mechanism Extraction Agent", "Extracts competitive and constraint mechanisms such as switching cost, network effect, regulatory headwind, platform dependency, and distribution control. Its uncertainty value feeds the final decision."],
        ["Agent 6", "Precedent Routing Agent", "Uses attribution consistency, vector retrieval, novelty detection, and deterministic route selection to decide whether the case is well-supported, conflicting, novel, or uncertain."],
        ["Agent 7", "Decision Agent", "Combines risk decomposition, conflict detection, offsetting analysis, and decision quality into the authoritative decision state."],
        ["Agent 8", "Synthesis and Interaction Agent", "Builds the strict response payload, writes strategic text with signal anchors, and keeps chat behavior grounded in the prior decision envelope."],
        ["Agent 9", "Monitoring and Feedback Agent", "Maintains prediction logs, drift checks, outcome logging, and calibration metrics without changing decisions automatically."],
    ]
    add_table(doc, ["Agent", "Role", "Summary"], agent_rows, [0.7, 1.55, 4.7])

    add_heading(doc, "3.0 Evidence Contract Across Agents", 2)
    add_para(doc, (
        "MIDAN is best understood as an evidence contract rather than a chain of text generators. "
        "Each agent receives bounded inputs, produces named artifacts, and hands forward only the evidence that the next agent is allowed to use. "
        "This contract is what keeps the system readable. It also makes error analysis possible because a bad decision can be traced to a gate, a market signal, a mechanism inference, a route, or a final decision rule."
    ))
    add_table(doc, ["Agent", "Primary Inputs", "Primary Outputs", "Failure Behavior"], [
        ["Agent 1", "Raw idea text.", "Validity envelope with rejection type, severity, missing information, and repair guidance.", "Stops inference before any model runs."],
        ["Agent 2", "Valid idea text and parsed sector/country hints.", "Confidence-scored schema and consistency result.", "Returns clarification-required state when required fields are unknown or inconsistent."],
        ["Agent 3", "Validated idea schema, sector, country, and loaded model artifacts.", "Macro vectors, regime, confidence, fuzzy membership, SHAP signals, trend, freshness, and routing score.", "Uses neutral or explicit skipped values for unavailable optional artifacts."],
        ["Agent 4", "Agent 2 schema and Agent 3 regime.", "Structured reasoning envelope for differentiation, competition, business model, unit economics, and interactions.", "Marks unavailable modules rather than inventing reasoning."],
        ["Agent 5", "Agent 2 values, Agent 4 reasoning, and idea text.", "Mechanism envelope, market structure, tensions, uncertainty, and epistemic summary.", "Returns insufficient-information envelope with high uncertainty."],
        ["Agent 6", "Market vector, SHAP shares, trend, retrieval artifacts, and routing score.", "Deterministic route path and human-review flag.", "Suppresses retrieval on novelty or unavailable artifacts."],
        ["Agent 7", "Risks, conflicts, mechanism uncertainty, route flags, and decision quality signals.", "Decision state, risk decomposition, conflict list, offsets, and reasoning trace.", "Blocks commitment under insufficient data or severe unresolved conflict."],
        ["Agent 8", "Complete decision envelope and strategic anchors.", "Strict response payload and grounded interaction behavior.", "Falls back to deterministic replies and sanitizes internal labels."],
        ["Agent 9", "Prediction and outcome records.", "Drift status, calibration metrics, and audit logs.", "Returns safe defaults when logs or baselines are unavailable."],
    ], [0.7, 1.55, 1.85, 2.5])

    add_heading(doc, "3.1 Agent 1: Intake and Validity", 2)
    for para in [
        "Agent 1 is the first control point. Its purpose is not to classify markets but to decide whether classification should be attempted at all. The implementation checks for contradictions, spam, prompt-injection patterns, non-viable claims, unsustainable economics, absence of value exchange, excessive vagueness, and insufficient length. A rejected input returns a structured envelope containing severity, rejection type, missing information, and suggested fixes.",
        "This design matters because startup-evaluation systems are vulnerable to false precision. A vague phrase such as an application that helps everyone should not be stretched into a market thesis. MIDAN therefore treats clarification as a valid system output. This is a research-relevant design choice: the absence of evidence is not converted into a weak prediction; it is preserved as a halt condition.",
    ]:
        add_para(doc, para)

    add_heading(doc, "3.2 Agent 2: Confidence-Scored Feature Extraction", 2)
    for para in [
        "Agent 2 turns unstructured idea text into a compact schema. The implementation extracts sector, country, business model, target segment, monetization type, stage, differentiation score, market readiness, regulatory risk, capital intensity, and competition intensity. Each field carries a value, confidence, and source. Required fields are checked before inference proceeds.",
        "The important feature is the unknown sentinel. Instead of filling missing fields with defaults and allowing downstream agents to treat them as observed, MIDAN labels them as unknown. Runtime defaults are used only after the sufficiency gate has passed and only where non-required fields need neutral arithmetic behavior. The distinction between extraction evidence and runtime fallback is preserved in the response.",
    ]:
        add_para(doc, para)

    add_heading(doc, "3.3 Agent 3: Market Intelligence", 2)
    doc.add_picture(str(figs[1]), width=Inches(7.2))
    add_caption(doc, "Figure 2. Internal flow of the Market Intelligence Agent.")
    for para in [
        "Agent 3 builds the market evidence envelope. It begins with static country and sector tables, then computes an effective macro vector using inflation, gross domestic product growth, macro friction, capital concentration, and sector velocity. Idea-derived adjustments may modify this vector, but only when the corresponding Agent 2 field exceeds the configured confidence floor. Each adjustment is returned with its source field, source value, affected feature, delta, and reason code.",
        "The adjusted vector is passed through a support vector machine classifier with rule-based overrides [1]. A fuzzy membership signal is then computed in the reduced market space, producing a top cluster, membership value, entropy, and ambiguity flag [3]. A LightGBM surrogate supplies per-request SHAP contributions [4, 5], while precomputed SARIMA forecasts provide a normalized sector-trend signal [6]. MIDAN also computes an intelligent score that combines regime favorability, SVM probability margin, fuzzy membership, SARIMA trend, and SHAP attribution consistency. The score is a routing signal only. It is not allowed to set the final decision.",
        "This agent is also responsible for transparency. The response exposes the base macro vector, adjusted macro vector, applied deltas, regime decision path, fuzzy membership, data freshness envelope, SHAP weights, and trend signal. The system therefore gives readers enough information to distinguish observed macro context, inferred idea effects, trained model output, and stale forecast risk.",
    ]:
        add_para(doc, para)
    add_table(doc, ["Signal", "Implementation Source", "Use in MIDAN", "Interpretation Boundary"], [
        ["Regime label", "SVM classifier with rule path.", "Defines the macro environment used by risk decomposition.", "It is a market-state label, not a startup success prediction."],
        ["SVM margin", "Difference between the highest and second-highest class probabilities.", "Contributes to the routing score.", "A wide margin means classifier separation, not independent evidence."],
        ["FCM membership", "Fuzzy membership around PCA-space centers.", "Identifies cluster fit and ambiguity.", "It is a soft geometry signal, not a second trained classifier."],
        ["SHAP weights", "Per-request LightGBM surrogate attribution.", "Explains which macro features shaped the classification.", "It explains the surrogate response, not causal market truth."],
        ["SARIMA trend", "Precomputed sector forecast table.", "Modulates routing and freshness-aware confidence.", "It is stale when not refreshed and is reported as such."],
        ["Intelligent Score", "Weighted routing composite.", "Selects the reasoning path used by Agent 6.", "It never sets the final decision state."],
    ], [1.2, 1.7, 1.8, 2.2])

    add_heading(doc, "3.4 Agent 4: Structured Reasoning", 2)
    for para in [
        "Agent 4 provides the idea-level interpretation surface. It evaluates differentiation, competition, business-model viability, unit economics, and signal interactions. The module does not introduce a new authoritative decision score. A legacy scalar is retained for backward compatibility, but the primary output is a structured reasoning envelope.",
        "The differentiation analyzer distinguishes sector-level baseline assumptions from idea-inferred evidence. The competition analyzer separates known competitor context from text-derived competitive pressure. The business-model analyzer emits money-flow and cost-structure information. The unit-economics analyzer produces qualitative proxies for customer acquisition pressure, revenue per user, and scalability. Signal interactions capture cases where evidence amplifies or conflicts, such as strong differentiation under high competitive pressure.",
        "The agent is intentionally modular. If a required field is unavailable, the affected analyzer declares insufficient information rather than manufacturing an interpretation. This behavior is important for later decision quality, because Agent 7 can distinguish a negative finding from an unavailable finding.",
    ]:
        add_para(doc, para)

    add_heading(doc, "3.5 Agent 5: Mechanism Extraction", 2)
    doc.add_picture(str(figs[2]), width=Inches(7.2))
    add_caption(doc, "Figure 3. Mechanism extraction flow from structural signals to uncertainty.")
    for para in [
        "Agent 5 extracts structural competitive mechanisms from the Agent 2 and Agent 4 evidence. Its mechanism vocabulary includes network effect, switching cost, brand moat, data moat, regulatory moat, cost advantage, process efficiency, distribution control, API dependency, platform dependency, and regulatory headwind. These are not decorative labels. Each mechanism carries confidence, evidence strength, inference depth, implication ceiling, and a reasoning trace.",
        "The implementation begins with an extractability gate. If structural signal coverage is too low, the agent returns an insufficient-information envelope with uncertainty set high. If extraction is possible, the pipeline translates idea text and reasoning fields into typed observations, assigns possible mechanisms, calibrates confidence, normalizes weights, derives market structure, classifies tensions, analyzes replication difficulty, applies contextual adjustments, propagates uncertainty, and builds an epistemic summary.",
        "The epistemic summary is central to publication-quality interpretation. It separates observed signals from inferred mechanisms and unresolved uncertainty. This prevents the synthesis agent from overstating speculative claims. Mechanism uncertainty is also passed to Agent 7 as a continuous probabilistic modifier, so weak mechanism evidence can affect decision quality without dictating the final verdict.",
    ]:
        add_para(doc, para)
    add_table(doc, ["Mechanism Family", "Mechanisms in the Implementation", "What the Agent Looks For"], [
        ["Advantage mechanisms", "Network effect, switching cost, brand moat, data moat, regulatory moat.", "Evidence that the startup can become harder to copy or displace as adoption grows."],
        ["Operational mechanisms", "Cost advantage and process efficiency.", "Evidence that the startup can perform the same job cheaper, faster, or with lower operational load."],
        ["Distribution mechanisms", "Distribution control.", "Evidence that access to customers, suppliers, or channels is structurally privileged."],
        ["Constraint mechanisms", "API dependency, platform dependency, regulatory headwind.", "Evidence that the idea may be constrained by external platforms, policy, or gatekeepers."],
    ], [1.5, 2.5, 2.8])
    add_para(doc, (
        "A useful detail is that Agent 5 does not treat every detected mechanism as equally strong. "
        "The extraction path separates directly observed evidence from one-step inference and from unsupported speculation. "
        "The implication ceiling then limits what the synthesis agent may say. "
        "For example, directly observed switching-cost evidence can support a stronger strategic conclusion than a generic claim about being innovative. "
        "This prevents mechanism labels from becoming decorative."
    ))

    add_heading(doc, "3.6 Agent 6: Precedent Routing", 2)
    for para in [
        "Agent 6 is a routing agent rather than a decision maker. It computes attribution consistency by comparing the current SHAP attribution vector against the mean attribution vector of the assigned cluster. When the artifact is unavailable or degenerate, the signal returns a neutral value. This prevents artifact absence from being misread as negative evidence.",
        "The explicit precedent lookup uses a vector composed of the scaled macro features and SHAP shares. The system performs nearest-neighbor retrieval over the available training-grid space and derives a majority-vote market regime, following the broader retrieval principle that relevant stored evidence can support reasoning when its scope is controlled [7, 13]. A novelty gate suppresses retrieval when the query is too far from known points. SARIMA trend can amplify or dampen confidence in precedent because historical neighbors are less informative when sector direction has shifted.",
        "The ReAct-style router maps intelligent score, attribution consistency, retrieval vote, novelty, and trend into deterministic paths. These paths include high certainty, low certainty, borderline confirmed, borderline conflict, atypical but supported, full conflict, maximum uncertainty, and novelty. The router can require human review, but it cannot set the final decision state. That authority remains with Agent 7.",
    ]:
        add_para(doc, para)
    add_table(doc, ["Route Family", "Interpretation", "Effect on Later Agents"], [
        ["Novelty", "The case sits outside useful precedent space, so retrieval is suppressed.", "Agent 7 still decides from macro and reasoning evidence, with novelty surfaced as framing."],
        ["High certainty", "The routing score is high and no serious attribution or retrieval conflict is active.", "Agent 7 may proceed without human-review override, subject to risk decomposition."],
        ["Low certainty", "The routing score is clearly unfavorable but attribution is reliable.", "Agent 7 treats the evidence as a coherent negative signal rather than ambiguity."],
        ["Borderline confirmed", "The score is borderline but retrieval agrees with the regime classification.", "Agent 7 receives moderate support rather than a forced commitment."],
        ["Borderline conflict", "The score is borderline and retrieval disagrees with the regime classification.", "A conflict is surfaced but does not automatically block the decision."],
        ["Full conflict", "Attribution is atypical and retrieval disagrees with the regime classification.", "The route forces human review and can produce INSUFFICIENT_DATA."],
        ["Maximum uncertainty", "No reliable second opinion is available or route evidence is too weak.", "The decision agent receives an explicit human-review requirement."],
    ], [1.35, 2.65, 2.9])

    add_heading(doc, "3.7 Agent 7: Decision Agent", 2)
    doc.add_picture(str(figs[3]), width=Inches(7.2))
    add_caption(doc, "Figure 4. Decision-state flow used by the Decision Agent.")
    for para in [
        "Agent 7 is the decision authority. It replaces a single-score verdict with a state machine grounded in risk decomposition, conflict detection, offsetting, and decision quality. Market risk is derived from the market regime, confidence, fuzzy ambiguity, and data freshness. Execution risk is derived from business-model viability and unit-economics proxies. Timing risk is derived from stage, readiness, and market regime alignment.",
        "Conflict detection is severity-ranked. Medium conflicts can raise caution without blocking a decision. High-severity unresolved conflicts can produce a conflicting-signals state. The agent also performs offsetting analysis. For example, structural differentiation may reduce market risk, and improved scalability can reduce execution risk, but offsets are logged explicitly and do not erase uncertainty.",
        "The final decision states include GO, CONDITIONAL, NO_GO, HIGH_UNCERTAINTY, CONFLICTING_SIGNALS, and INSUFFICIENT_DATA. The decision strength is qualitative. The legacy numerical score is retained only for compatibility and is explicitly marked as having no decision authority.",
    ]:
        add_para(doc, para)
    add_table(doc, ["Decision State", "Meaning", "Typical Trigger"], [
        ["GO", "The evidence supports forward motion under the current assumptions.", "No blocking conflict and risk levels remain manageable."],
        ["CONDITIONAL", "The idea is promising but requires targeted validation before commitment.", "One or more material assumptions remain unresolved."],
        ["NO_GO", "The system does not support building or committing capital under the current framing.", "Risk remains dominant after offsetting analysis."],
        ["HIGH_UNCERTAINTY", "The system can advise but should not issue a firm recommendation.", "Decision quality is weak or signal agreement is low."],
        ["CONFLICTING_SIGNALS", "The evidence contains unresolved severe contradictions.", "High-severity conflict requires resolution before decision."],
        ["INSUFFICIENT_DATA", "The system lacks enough reliable evidence for a decision.", "Missing required reasoning fields or human-review override."],
    ], [1.3, 2.5, 3.0])
    add_para(doc, (
        "The decision-state vocabulary is deliberately small. "
        "A larger set of states would make the response feel richer but would weaken the interpretability of the system. "
        "The implemented states cover the practical outcomes needed by an evaluator: proceed, validate, stop, advise cautiously, resolve conflict, or collect more information."
    ))

    add_heading(doc, "3.8 Agent 8: Synthesis and Interaction", 2)
    for para in [
        "Agent 8 converts the evidence envelope into the user-facing response. It builds a strict Pydantic response payload, adds strategic text fields, attaches signal anchors to each major text field, and sanitizes chat output so internal labels and unknown sentinels do not leak into user-facing conversation.",
        "The interaction behavior is post-decision aware. A normal decision enters standard-advisor mode. A conflicting-signals state enters conflict-resolution mode. A high-uncertainty state becomes advisory only. An insufficient-data state asks for clarification rather than providing confident advice. This keeps conversation aligned with the decision evidence instead of allowing a chat model to improvise beyond the system state.",
    ]:
        add_para(doc, para)

    add_heading(doc, "3.9 Agent 9: Monitoring and Feedback", 2)
    for para in [
        "Agent 9 records prediction signals and supports drift and calibration workflows. The prediction log stores timestamp, sector, country, regime, SVM margin, intelligent score, SHAP consistency, fuzzy membership, SARIMA trend, route path, and decision state. Drift detection uses a dual-signal gate: classification-margin decline and fuzzy-centroid displacement must both fire before drift is confirmed.",
        "Outcome feedback is separate. The outcome module logs externally observed outcomes using an append-only record that stores decision identifier, predicted state, regime, sector, country, route path, intelligent score, and outcome value. It does not store idea text. Calibration metrics remain advisory and do not automatically retrain or alter routing thresholds.",
    ]:
        add_para(doc, para)

    doc.add_page_break()
    add_heading(doc, "4. Implementation", 1)
    doc.add_picture(str(figs[4]), width=Inches(7.2))
    add_caption(doc, "Figure 5. Supporting data acquisition subsystem.")
    impl = [
        "The implementation is organized as a Python package with explicit ownership boundaries. The FastAPI application is defined in the endpoint module, while the root API file remains a compatibility shim for older imports. The core module owns request schemas, loaded artifacts, reference tables, utility functions, data freshness logic, and strict response models. The orchestration module wires the agents together.",
        "The system uses scikit-learn for scaling, PCA, and support-vector classification [8]; LightGBM as a surrogate model for explainability [4]; SHAP for feature-attribution signals [5]; statsmodels for SARIMA modeling [9]; scikit-fuzzy for fuzzy membership [3]; Pydantic for schema enforcement [15]; FastAPI for HTTP endpoints [14]; Streamlit and a static HTML frontend for interaction surfaces [20]; and Gemini-backed language-model calls for selected synthesis and conversation paths when configured. When language-model paths fail or are unavailable, the implementation falls back to deterministic heuristics rather than failing silently.",
        "Model artifacts are loaded at startup from the project model directory. Required artifacts include the scaler, PCA model, SVM classifier, label encoder, LightGBM surrogate, SARIMA result table, competitor context, sentiment context, and fuzzy centers. Several newer artifacts are optional and fail soft. If SHAP cluster means are absent, attribution consistency returns a neutral value. If the retrieval index is unavailable, precedent retrieval is skipped with a reason code. If the drift baseline is unavailable, drift detection returns a safe no-baseline result.",
        "The runtime pipeline begins with input validation, then performs sector and country parsing, confidence-scored idea extraction, and sufficiency checks. Only after these checks pass does the system run macro intelligence, structured reasoning, mechanism extraction, routing, decision derivation, synthesis, logging, and response construction. This ordering is important because it prevents invalid ideas from entering model inference and prevents missing idea fields from contaminating macro or decision outputs.",
        "The data acquisition subsystem is implemented separately under the project-scraped-data directory. It collects raw startup material from websites, Failory, YC, Reddit, and Product Hunt where configured. The pipeline classifies sources, runs extractor modules, validates structured entries, resolves conflicts across sources, analyzes patterns, applies a decision-analysis pass, and writes a structured startup-intelligence JSON artifact. A local LLM supplement is optional and is constrained to filling empty fields in ambiguous narrative cases. It does not overwrite rule-based extraction results.",
        "The response contract is intentionally strict. The analyze, project, and interact endpoints return a common response payload containing decision state, decision strength, decision quality, risk decomposition, reasoning trace, post-decision mode, and supporting evidence. Malformed payloads raise validation errors at the boundary rather than being silently corrected. This gives the system a stable public contract and allows downstream interfaces to treat missing data as explicit rather than accidental.",
        "The implementation also includes reliability mechanisms: structured request identifiers, failure logs, prediction logs, drift checks, outcome logging, and schema tests. The project does not use external monitoring, persistence infrastructure, or complex configuration frameworks. Configuration is centralized in a plain Python module, which keeps thresholds and feature toggles inspectable without adding operational machinery that would obscure the research prototype.",
    ]
    for para in impl:
        add_para(doc, para)
    add_heading(doc, "4.1 Runtime APIs and Artifacts", 2)
    add_table(doc, ["Interface or Artifact", "Role in the System", "Research Relevance"], [
        ["/analyze", "Runs the full pipeline and returns the strict response payload.", "Primary path for evaluating a structured startup idea."],
        ["/project", "Handles projection-style analysis, probe questions, and low-quality input branches.", "Shows that pre-analysis and decided states share one schema."],
        ["/interact", "Classifies conversational intent before deciding whether to run analysis.", "Prevents casual or partial messages from forcing false analysis."],
        ["/chat", "Provides post-decision conversation grounded in the prior decision envelope.", "Keeps follow-up advice aligned with the evidence state."],
        ["Prediction log", "Stores regime, route, score components, and decision state.", "Provides the basis for drift monitoring and runtime audits."],
        ["Outcome log", "Stores externally submitted outcomes without idea text.", "Enables long-term calibration without leaking raw startup descriptions."],
    ], [1.3, 2.6, 3.0])
    add_para(doc, (
        "The API structure is important because it separates analysis from conversation. "
        "A common failure in advisory systems is that the chat surface begins to behave as the decision engine. "
        "MIDAN avoids this by storing the decision envelope and routing follow-up turns according to the existing decision state. "
        "The chat system may explain, clarify, or ask for missing information, but it should not invent a new decision path independent of Agent 7."
    ))
    add_heading(doc, "4.2 Data and Artifact Provenance", 2)
    add_para(doc, (
        "The project contains five raw CSV datasets used or referenced by the training and context-building workflow: a financial-news sentiment dataset, a startup success dataset, venture-investment records, unicorn startup records, and World Bank macro indicators [11]. "
        "The notebook then builds feature vectors, model artifacts, and context JSON files. "
        "The research chapter intentionally avoids reporting file sizes because they are not meaningful scientific evidence and can distract from the structure of the system."
    ))
    add_table(doc, ["Data or Artifact Type", "Observed Local Form", "How It Is Used"], [
        ["Startup and funding records", "CSV datasets and structured startup JSON.", "Support training, context, and startup-intelligence extraction."],
        ["Macro indicators", "World Bank-derived CSV plus static runtime tables.", "Inform country and market-condition features."],
        ["Classifier artifacts", "Scaler, PCA, SVM, label encoder, and LightGBM surrogate.", "Classify market regime and produce explainability signals."],
        ["Time-series artifacts", "SARIMA result table and sector models.", "Provide sector trend signals and staleness checks."],
        ["Context artifacts", "Competitor and sentiment JSON files.", "Support synthesis and contextual interpretation."],
        ["Optional routing artifacts", "SHAP cluster means, retrieval index, labels, and drift baseline when generated.", "Improve routing when present and degrade safely when absent."],
    ], [1.6, 2.1, 3.1])
    add_heading(doc, "4.3 Construction Logic", 2)
    for para in [
        "The construction logic follows a fail-soft pattern. Core model artifacts are required for full analysis, while newer routing artifacts are optional and explicitly degraded when absent. This lets the system continue operating in a reduced-evidence mode without pretending that all evidence sources are active.",
        "The implementation also separates training-time artifacts from runtime artifacts. DBSCAN and certain PCA visualizations are retained for reproducibility of the notebook and for understanding the market-space construction, but they are not loaded as runtime decision components. This distinction is useful in the chapter because it prevents readers from assuming that every artifact in the model directory participates in live inference.",
        "The pipeline was also designed to preserve backwards compatibility. Legacy fields such as the numerical score and decision badge still appear for older consumers, but the response text and decision logic emphasize the newer decision state, decision strength, and risk decomposition. This is a common transitional problem in applied systems research: a better decision interface must be introduced without breaking existing frontends or tests."
    ]:
        add_para(doc, para)
    add_table(doc, ["Implementation Principle", "Concrete Form in MIDAN"], [
        ["Fail fast on bad inputs", "Agent 1 and Agent 2 stop inference before market models run."],
        ["Degrade safely", "Optional retrieval and attribution artifacts return neutral or skipped states when unavailable."],
        ["Separate authority", "Routing score influences path selection but does not set the decision state."],
        ["Expose provenance", "Base macro values, inferred deltas, regime path, and freshness metadata are returned."],
        ["Keep schema strict", "The public payload validates required decision, quality, risk, and trace fields."],
        ["Preserve compatibility", "Legacy score fields remain visible but are marked as non-authoritative."],
    ], [2.0, 4.7])

    doc.add_page_break()
    add_heading(doc, "5. Results", 1)
    res_intro = [
        "The results reported here are limited to real signals found or produced during inspection of the project folder. They should be read as system-behavior results rather than external investment-performance claims. No synthetic benchmark outcomes are invented.",
        "The implemented Phase 5 validation script passed 39 of 39 checks. These checks cover intelligent-score construction, correlated-signal discounting, missing-artifact behavior, attribution consistency defaults, novelty handling, deterministic route selection, drift-gate behavior, and the integration between routing and the decision state machine. The validation run also confirmed that degraded optional artifacts do not crash the pipeline.",
    ]
    for para in res_intro:
        add_para(doc, para)
    add_table(doc, ["Validation Area", "Observed Result"], [
        ["Intelligent Score", "6 checks passed, including weight normalization, neutral unknown-regime handling, and SVM margin computation."],
        ["Routing and RAG Degradation", "Attribution and retrieval degraded safely when artifacts were absent. FAISS-specific tests were skipped because FAISS was not installed in the current runtime."],
        ["Router", "10 routing checks passed across novelty, high certainty, low certainty, borderline confirmation, conflict, full conflict, and maximum uncertainty."],
        ["Drift Monitor", "6 checks passed, including dual-signal confirmation, insufficient-log handling, and log trimming."],
        ["Decision Integration", "6 checks passed, including human-review override to INSUFFICIENT_DATA and RAG conflict injection."],
        ["Integration", "3 checks passed, including clean pipeline import and degraded-artifact operation."],
    ], [1.6, 5.2])

    add_heading(doc, "5.1 Scenario 1: Structured SaaS Input", 2)
    for para in [
        "A structured restaurant SaaS idea for Cairo was processed successfully. The system identified the case as an emerging-market regime and selected the high-certainty routing path. The final decision state was GO, with an intelligent score of 0.785 and Agent 2 aggregate confidence of 0.933. Risk decomposition produced medium market risk, low execution risk, and medium timing risk.",
        "This scenario illustrates the intended full-pipeline behavior. The input passed validity and feature sufficiency checks, moved through market intelligence and structured reasoning, and reached a decision state without requiring clarification or human review. The result should not be interpreted as a real investment recommendation; it is an observed behavior of the implemented pipeline under a well-formed input.",
    ]:
        add_para(doc, para)

    add_heading(doc, "5.2 Scenario 2: Marketplace Under Higher Friction", 2)
    for para in [
        "A Saudi last-mile logistics marketplace input also passed parsing, with Agent 2 aggregate confidence of 0.900. The system classified the market context as HIGH_FRICTION_MARKET and selected the maximum-uncertainty route. The final state was INSUFFICIENT_DATA. Market, execution, and timing risks were all high.",
        "This scenario demonstrates a useful distinction: the input was parseable, but the evidence path did not support an automated commitment decision. The router required review, and the decision agent preserved that uncertainty as a blocking state. This behavior is consistent with the architecture, where high confidence in parsing does not imply high confidence in strategic action.",
    ]:
        add_para(doc, para)

    add_heading(doc, "5.3 Scenario 3: Under-Defined Input", 2)
    for para in [
        "A vague input stating that an application uses AI to make life better was rejected before market inference. The system returned a clarification-required style rejection with the reason vague_non_actionable. No regime, intelligent score, route path, or risk decomposition was produced.",
        "This scenario validates the fail-fast principle. A weakly specified idea did not receive a manufactured market analysis. Instead, the system preserved the true state of evidence: the user must define the problem, customer, and mechanism before a decision can be formed.",
    ]:
        add_para(doc, para)

    add_table(doc, ["Scenario", "Decision State", "Regime", "Route", "Key Metrics"], [
        ["Restaurant SaaS in Cairo", "GO", "EMERGING_MARKET", "PATH_1_HIGH_CERTAINTY", "IS 0.785; Agent 2 confidence 0.933; risks medium/low/medium."],
        ["Saudi logistics marketplace", "INSUFFICIENT_DATA", "HIGH_FRICTION_MARKET", "PATH_7_MAXIMUM_UNCERTAINTY", "IS 0.666; Agent 2 confidence 0.900; risks high/high/high."],
        ["Generic AI app", "Clarification required", "Not run", "Not run", "Rejected as vague_non_actionable before market inference."],
    ], [1.5, 1.35, 1.35, 1.6, 2.1])

    add_heading(doc, "5.4 Runtime Log Observations", 2)
    for para in [
        "The local prediction log contained 17 entries at inspection time. All logged entries were SaaS cases. Fourteen were associated with Egypt and three with Saudi Arabia. All were classified as EMERGING_MARKET. Fifteen used the maximum-uncertainty route and two used the high-certainty route. Decision states in the log included 15 INSUFFICIENT_DATA outcomes, one GO outcome, and one CONDITIONAL outcome.",
        "These log observations are useful because they show the system preserving uncertainty frequently under the tested local workload. They are not enough to estimate external predictive validity. They also show the need for broader scenario coverage across sectors, countries, and business models before presenting aggregate performance claims.",
    ]:
        add_para(doc, para)

    add_heading(doc, "5.5 Data Acquisition Output", 2)
    for para in [
        "The supporting data acquisition subsystem produced a structured startup-intelligence artifact with 46 startup entries. Its metadata identifies the pipeline version as 2.0.0-phase2 and lists website and insight sources. The training signal corpus contains 20 entries, the validation ground-truth file contains 30 entries, and the extraction error log contains 18 cases.",
        "These artifacts demonstrate that the project includes a data pipeline beyond the runtime inference path. However, they should be treated as supporting evidence for system construction rather than a complete empirical benchmark. The current chapter therefore reports them as implementation outputs, not as proof of investment accuracy.",
    ]:
        add_para(doc, para)
    add_heading(doc, "5.6 What the Results Do and Do Not Prove", 2)
    for para in [
        "The validation results prove that several internal contracts behave as designed in the inspected environment. The intelligent score computes its components correctly. The router maps signals to deterministic paths. Drift detection requires both monitored signals before confirming drift. The decision agent accepts route-level human-review signals and converts them into an insufficient-data state when appropriate.",
        "The results do not prove investment accuracy. They do not show that a GO decision predicts startup success, that a NO_GO decision predicts failure, or that the mechanism extractor has high recall against a broad expert-labeled corpus. Those claims require longitudinal outcome records, held-out scenario sets, and comparison against expert baselines. The current evidence is best described as implementation validation and behavioral observation.",
        "This distinction is important for a research chapter. A strong architecture can still require empirical calibration. MIDAN should therefore be presented as an implemented, traceable decision-support architecture with promising internal validation, not as a fully validated venture-selection model."
    ]:
        add_para(doc, para)

    doc.add_page_break()
    add_heading(doc, "6. Limitations and Challenges", 1)
    for para in [
        "MIDAN is a sophisticated research prototype, but several limitations remain. First, the macro data used by the runtime is not live. The system exposes this explicitly through the freshness envelope and the validation script reported a macro staleness alert. This is valuable transparency, but it also means that deployed use would require a reliable refresh process for macro tables and sector forecasts.",
        "Second, the logged runtime evidence is narrow. The local prediction log observed during inspection covers SaaS cases only, with a small country distribution. This prevents credible claims about general performance across fintech, ecommerce, healthtech, edtech, logistics, agritech, and other sectors. Broader evaluation must include balanced scenario coverage and real outcome follow-up.",
        "Third, optional retrieval artifacts are not always available in the current runtime. The system degrades safely when attribution or retrieval artifacts are absent, and the validation script confirms this behavior. Safe degradation is a strength, but it also reduces the amount of evidence available to the routing agent. Future evaluation should distinguish full-artifact runs from degraded runs.",
        "Fourth, mechanism extraction is deterministic and evidence-calibrated, but it remains dependent on the quality of upstream idea text and structured reasoning fields. Weak or vague inputs correctly reduce extraction quality. The challenge is to strengthen mechanism extraction without overclaiming from thin text.",
        "Fifth, the decision state machine is rule-governed. This improves traceability, but rule systems require careful calibration. The current risk thresholds, offsetting rules, and conflict severities are defensible design choices, not yet empirically optimized decision policies. A future calibration study should compare these states against longitudinal outcomes.",
        "Sixth, the system can use language-model calls for extraction, synthesis, or conversation when configured. These paths are bounded by fallback logic and response sanitization, but publication-quality deployment still requires systematic audits for stability, hallucination resistance, and consistency under repeated prompts.",
        "Finally, the system currently evaluates startup concepts, not complete companies. A real investment decision would require founder history, financial model, customer evidence, legal exposure, fundraising context, and live competitive intelligence. MIDAN should therefore be viewed as a structured decision-support system rather than a replacement for expert due diligence.",
    ]:
        add_para(doc, para)
    add_table(doc, ["Challenge", "Why It Matters", "Current Mitigation"], [
        ["Data staleness", "Market signals decay quickly, especially for inflation, funding conditions, and sector momentum.", "Freshness envelope and staleness penalty expose the issue at runtime."],
        ["Limited logged scenarios", "Narrow logs cannot support broad empirical claims.", "Results section treats logs as observations, not benchmarks."],
        ["Optional artifact availability", "Routing evidence weakens when retrieval or attribution artifacts are absent.", "Neutral defaults and explicit skipped reasons avoid false certainty."],
        ["Rule calibration", "Risk and offset rules may encode reasonable priors but still need empirical tuning.", "Rules are centralized and traceable, making future calibration feasible."],
        ["Language-model variability", "LLM paths can introduce instability if unconstrained.", "Fallback paths, prompt contracts, and sanitization reduce exposure."],
        ["Outcome scarcity", "Without longitudinal outcomes, decision quality cannot be externally measured.", "Outcome logging and calibration metrics are already implemented."],
    ], [1.5, 2.6, 2.8])

    doc.add_page_break()
    add_heading(doc, "7. Conclusion + Future Work", 1)
    for para in [
        "MIDAN demonstrates how a startup-evaluation system can combine machine learning, structured reasoning, mechanism extraction, and calibrated decision synthesis without collapsing the result into a single opaque score. Its central contribution is architectural discipline. Each agent has a bounded role, each uncertainty source is surfaced, and the final decision belongs to a state machine that can halt, advise, or commit depending on the evidence.",
        "The inspected implementation supports the major claims of the architecture. It includes confidence-aware parsing, transparent macro adjustment, regime classification, fuzzy membership, SHAP explainability, sector trend signals, routing logic, mechanism analysis, strict response schemas, chat grounding, prediction logging, drift checks, and outcome-feedback infrastructure. The Phase 5 validation run passed all 39 checks in the available environment. Local scenarios also showed differentiated behavior across well-formed, uncertain, and under-defined inputs.",
        "The most important conclusion is not that MIDAN is already a complete empirical evaluator. It is that the system establishes a defensible structure for evaluation: block bad inputs, separate evidence types, avoid hidden defaults, preserve uncertainty, and make recommendations only at the level the evidence can support. That structure is a strong basis for future empirical work.",
    ]:
        add_para(doc, para)

    add_heading(doc, "7.1 Future Work", 2)
    future_items = [
        "Larger datasets: expand the macro, startup, funding, and failure corpora across more sectors, geographies, and time periods.",
        "Real-world validation: compare MIDAN decisions with expert evaluations and subsequent startup outcomes rather than relying only on internal consistency checks.",
        "Longitudinal outcome tracking: increase use of the outcome feedback module so decision states can be calibrated against observed validation, invalidation, or partial outcomes over time.",
        "Adaptive macro updates: replace static macro tables with governed refresh jobs and provenance tracking for each update.",
        "Stronger mechanism extraction: expand evidence sources for competitive mechanisms, including public product pages, customer reviews, regulatory filings, and structured competitor data.",
        "Scaling infrastructure: move from local logs to durable storage and controlled background jobs once the prototype enters sustained use.",
        "Expanded country support: add more countries and regional macro profiles while documenting coverage and freshness for each market.",
        "Empirical calibration of Agent 4 reasoning: tune business-model, competition, and unit-economics rules against observed outcomes rather than only expert priors.",
        "Hybrid symbolic and language-model improvements: use symbolic guards for authority and consistency while allowing language models to help extract richer mechanisms under strict evidence constraints.",
    ]
    add_bullets(doc, future_items)
    add_table(doc, ["Future Work Track", "Research Question", "Expected Evidence"], [
        ["Dataset expansion", "Does performance remain stable across broader sectors and countries?", "Balanced scenario sets, updated macro data, and larger structured startup corpora."],
        ["Outcome validation", "Do decision states correlate with later validation or invalidation?", "Logged outcomes over defined horizons and per-regime calibration statistics."],
        ["Mechanism extraction", "Which mechanisms are reliably detectable from public startup evidence?", "Expert-labeled mechanism corpora and agreement analysis."],
        ["Macro adaptation", "How often should country and sector signals refresh?", "Staleness experiments comparing old and refreshed macro tables."],
        ["Hybrid reasoning", "Where should symbolic rules end and language-model extraction begin?", "Ablation studies comparing deterministic, LLM-assisted, and hybrid variants."],
        ["Infrastructure scaling", "What operational architecture preserves auditability at higher volume?", "Durable logs, background jobs, versioned artifacts, and reproducible deployment runs."],
    ], [1.5, 2.65, 2.75])
    add_para(doc, (
        "A mature version of MIDAN would therefore combine three feedback loops. "
        "The first is a data-refresh loop that keeps macro and sector signals current. "
        "The second is an outcome loop that compares decisions with observed venture progress. "
        "The third is a calibration loop that revises thresholds, risk offsets, and mechanism weights only after enough outcome evidence accumulates. "
        "Keeping these loops separate is essential. A system that retrains too quickly on sparse feedback can become less reliable, not more intelligent."
    ))
    add_para(doc, (
        "The system's long-term research value lies in that separation. "
        "MIDAN can be studied as an architecture for disciplined startup reasoning, as a testbed for hybrid symbolic and statistical decision systems, and as a practical interface for founders and analysts who need structured feedback without losing sight of uncertainty. "
        "The next stage is not to add more theatrical intelligence, but to collect enough real outcomes to calibrate the intelligence that is already there."
    ))

    add_heading(doc, "8. References", 1)
    refs = [
        "Cortes, C., Vapnik, V.: Support-vector networks. Machine Learning 20, 273-297 (1995). https://doi.org/10.1007/BF00994018",
        "Ester, M., Kriegel, H.-P., Sander, J., Xu, X.: A density-based algorithm for discovering clusters in large spatial databases with noise. Proceedings of the Second International Conference on Knowledge Discovery and Data Mining, 226-231 (1996).",
        "Bezdek, J.C.: Pattern Recognition with Fuzzy Objective Function Algorithms. Springer, Boston (1981). https://doi.org/10.1007/978-1-4757-0450-1",
        "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., Liu, T.-Y.: LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems 30 (2017).",
        "Lundberg, S.M., Lee, S.-I.: A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems 30 (2017).",
        "Box, G.E.P., Jenkins, G.M., Reinsel, G.C., Ljung, G.M.: Time Series Analysis: Forecasting and Control, 5th edn. Wiley (2015).",
        "Johnson, J., Douze, M., Jegou, H.: Billion-scale similarity search with GPUs. IEEE Transactions on Big Data 7(3), 535-547 (2021). https://doi.org/10.1109/TBDATA.2019.2921572",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., et al.: Scikit-learn: Machine learning in Python. Journal of Machine Learning Research 12, 2825-2830 (2011).",
        "Seabold, S., Perktold, J.: Statsmodels: Econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference, 92-96 (2010). https://doi.org/10.25080/Majora-92bf1922-011",
        "McKinney, W.: Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 56-61 (2010). https://doi.org/10.25080/Majora-92bf1922-00a",
        "World Bank: World Development Indicators. The World Bank Group (accessed 2026). https://databank.worldbank.org/source/world-development-indicators",
        "Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., Cao, Y.: ReAct: Synergizing reasoning and acting in language models. International Conference on Learning Representations (2023).",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., et al.: Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems 33, 9459-9474 (2020).",
        "FastAPI: FastAPI documentation. https://fastapi.tiangolo.com/",
        "Pydantic: Pydantic documentation. https://docs.pydantic.dev/",
        "Harris, C.R., Millman, K.J., van der Walt, S.J., Gommers, R., Virtanen, P., Cournapeau, D., et al.: Array programming with NumPy. Nature 585, 357-362 (2020). https://doi.org/10.1038/s41586-020-2649-2",
        "Virtanen, P., Gommers, R., Oliphant, T.E., Haberland, M., Reddy, T., Cournapeau, D., et al.: SciPy 1.0: Fundamental algorithms for scientific computing in Python. Nature Methods 17, 261-272 (2020). https://doi.org/10.1038/s41592-019-0686-2",
        "Hunter, J.D.: Matplotlib: A 2D graphics environment. Computing in Science and Engineering 9(3), 90-95 (2007). https://doi.org/10.1109/MCSE.2007.55",
        "Kluyver, T., Ragan-Kelley, B., Perez, F., Granger, B., Bussonnier, M., Frederic, J., et al.: Jupyter Notebooks, a publishing format for reproducible computational workflows. In: Positioning and Power in Academic Publishing: Players, Agents and Agendas, pp. 87-90 (2016). https://doi.org/10.3233/978-1-61499-649-1-87",
        "Streamlit: Streamlit documentation. https://docs.streamlit.io/",
        "Plotly Technologies Inc.: Collaborative data science. Montreal (2015). https://plotly.com/",
        "Uvicorn: Uvicorn documentation. https://www.uvicorn.org/",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"[{i}] {ref}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    build_doc()
    print("MIDAN_Research_Chapter_Revised.docx")
