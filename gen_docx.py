from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4

for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

def title(text, size=28, color=RGBColor(15,15,35)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.bold = True
    return p

def subtitle(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = RGBColor(100,116,139)
    return p

def heading(num, name, layer):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{layer}\n')
    r1.font.size = Pt(9); r1.font.color.rgb = RGBColor(59,130,246)
    r2 = p.add_run(f'{num}. {name}')
    r2.font.size = Pt(16); r2.bold = True; r2.font.color.rgb = RGBColor(30,41,59)
    p.paragraph_format.space_before = Pt(20)
    return p

def formula(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(10); r.font.color.rgb = RGBColor(15,23,42)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def label(text, color=RGBColor(29,78,216)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10); r.bold = True; r.font.color.rgb = color
    r.font.all_caps = True
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    return t

# ══════════ COVER ══════════
doc.add_paragraph()
doc.add_paragraph()
title('MIDAN Intelligence Engine')
subtitle('Complete Scoring Formulas Reference')
subtitle('Pipeline Version 1.0 — All formulas in execution order', 12)
doc.add_paragraph()
subtitle('Intelligent Systems Project — Faculty of Computer Science', 11)
doc.add_page_break()

# ══════════ TOC ══════════
p = doc.add_paragraph()
r = p.add_run('Table of Contents'); r.bold = True; r.font.size = Pt(16)
toc_items = [
    '1. Logical Validity Score (L0 Gate)',
    '2. L1 Aggregate Confidence (L1 Gate)',
    '3. Effective Inflation (L2A Macro)',
    '4. Macro Friction (L2A Macro)',
    '5. SARIMA Trend Normalization (L2D)',
    '6. FCM Fuzzy Membership (L2B.5)',
    '7. SHAP Normalized Shares (L2C)',
    '8. SHAP Cosine Similarity (L2C.5)',
    '9. Novelty Score (RAG)',
    '10. RAG Confidence + ARIMA Modifier (RAG)',
    '11. XAI Score (L2C)',
    '12. Idea Signal (L3 Scoring)',
    '13. Intelligent Score — IS (L2D.5 Routing)',
    '14. Legacy TAS (L4a Display Only)',
    '15. Decision Quality Assessment (L4)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Pipeline Flow: ')
r.bold = True
p.add_run('User Input → L0 Gate (validity) → L1 Parser (extraction) → L2 Macro (SVM, FCM, SHAP, SARIMA) → IS Composite (5-signal routing) → ReAct Router (7-path) → L3 Reasoning (idea signal) → L4 Decision Engine (GO / CONDITIONAL / NO_GO)')

doc.add_page_break()

# ══════════ FORMULA 1 ══════════
heading(1, 'Logical Validity Score', 'Layer 0 — Input Gate  |  GATE')
label('Formula')
formula('logical_validity_score ∈ [0.0, 1.0]  — fixed constant per rejection type')
add_table(['Rejection Type', 'Score', 'Confidence'], [
    ['adversarial_prompt', '0.02', '0.97'],
    ['logical_impossibility', '0.02', '0.97'],
    ['no_revenue_model', '0.04', '0.96'],
    ['unsustainable_economics', '0.05', '0.93'],
    ['spam_or_gibberish', '0.05–0.08', '0.90–0.95'],
    ['no_value_exchange', '0.12', '0.90'],
    ['contradictory_claims', '0.18', '0.92'],
    ['too_short', '0.20', '0.97'],
    ['vague_non_actionable', '0.35', '0.82'],
    ['LLM arbiter rejection', '1.0 − conf', '[0.85, 0.96]'],
    ['Pass (valid idea)', '0.92', '—'],
])
label('What It Measures')
body('How close the input is to being a real business concept. 0.0 = definitively not a business. 1.0 = definitively a real business.')
label('Why This Design', RGBColor(29,78,216))
body('Scores are hand-calibrated to reflect severity ordering: physically impossible ideas (0.02) are worse than vague ones (0.35). This is a gate, not a model — scores are ordinal labels that rank rejection severity. The LLM arbiter score is 1 − confidence because higher LLM confidence in rejection means lower validity. Pass-through is 0.92 (not 1.0) because passing L0 doesn\'t guarantee viability.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('A gate needs correct ordering (impossible < no revenue < vague < valid), not statistical calibration. The scores enforce that ordering by construction.')

# ══════════ FORMULA 2 ══════════
heading(2, 'Aggregate Confidence', 'Layer 1 — Parser  |  GATE')
label('Formula')
formula('aggregate_confidence = mean( confidence[f]  for f in required_fields )\n\nrequired_fields = {business_model, target_segment, stage, differentiation_score}\nconfidence[f] ∈ [0.0, 1.0]\n\nGate: halts pipeline if aggregate_confidence < 0.50\nField becomes UNKNOWN if individual confidence < 0.55')
label('What It Measures')
body('Sufficiency gate ensuring the LLM extracted required fields with enough certainty. Below threshold → pipeline asks for clarification.')
label('Why This Design', RGBColor(29,78,216))
body('Arithmetic mean treats all fields equally. We gate on the mean rather than any single field because one low-confidence field shouldn\'t block if the others are strong — L2–L4 can compensate for one uncertain input but not systemic ambiguity.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('The mean penalizes evenly when multiple fields are weak but allows one borderline field through when the rest are high.')

# ══════════ FORMULA 3 ══════════
heading(3, 'Effective Inflation', 'Layer 2A — Macro Vector  |  SIGNAL')
label('Formula')
formula('scale = base_inflation / 33.9\neffective_inflation = clip( eff_inf_offset × scale,  1.0,  100.0 )\n\nbase_inflation  = country inflation rate (COUNTRY_MACRO_DEFAULTS)\neff_inf_offset  = sector sensitivity (SECTOR_EFF_MACRO)\n33.9            = Egypt\'s inflation (reference denominator)')
label('What It Measures')
body('Converts raw country inflation into a sector-adjusted effective inflation reflecting how much inflation impacts this particular sector.')
label('Why This Design', RGBColor(29,78,216))
body('Fintech (offset 7.5) has less inflation pass-through than ecommerce (offset 36.0) — digital services have lower COGS exposure. Dividing by 33.9 normalizes so the SVM sees Egypt-calibrated values regardless of country.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Proportional rescaling: new = reference × (country/reference_country). Preserves relative ordering of sectors across countries while keeping magnitude in SVM training range.')

# ══════════ FORMULA 4 ══════════
heading(4, 'Macro Friction', 'Layer 2A — Macro Vector  |  SIGNAL')
label('Formula')
formula('macro_friction = clip( effective_inflation + unemployment − gdp_growth,  −50,  100 )')
label('What It Measures')
body('Single scalar capturing how much the macro environment resists startup creation. Higher = harder to operate.')
label('Why This Design', RGBColor(29,78,216))
bullet('Inflation adds friction — higher costs erode margins')
bullet('Unemployment adds friction — weaker consumer spending, harder hiring')
bullet('GDP growth reduces friction — expanding economy creates demand')
body('The subtraction makes it a net resistance score: a growing economy partially offsets inflation and unemployment drag.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Signed linear combination with clear directional interpretation. A proxy designed to separate favorable from hostile environments — exactly what the SVM needs as a feature input.')

# ══════════ FORMULA 5 ══════════
heading(5, 'SARIMA Trend Normalization', 'Layer 2D — Time Series  |  SIGNAL')
label('Formula')
formula('forecast_mean = mean( max(0, v)  for v in sarima_forecast_values )\nsarima_trend = clip( forecast_mean / 50.0,  0.15,  0.90 )\n\nDefault = 0.50 (neutral) when no SARIMA model exists')
label('What It Measures')
body('Normalized sector trajectory in [0.15, 0.90]. Higher = sector growing, lower = sector declining.')
label('Why This Design', RGBColor(29,78,216))
bullet('max(0, v) floors negative forecasts — negative deal counts are nonsensical')
bullet('Dividing by 50.0 maps to 0–1 range (50 deals/period = neutral midpoint)')
bullet('Clip to [0.15, 0.90] prevents extremes')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Min-max normalization with domain-informed midpoint. Divisor of 50 calibrated from training data. Asymmetric clip ensures no signal saturation.')

doc.add_page_break()

# ══════════ FORMULA 6 ══════════
heading(6, 'FCM Fuzzy Membership', 'Layer 2B.5 — Fuzzy Clustering  |  SIGNAL')
label('Formula')
formula('dᵢ = ‖x_pca − centerᵢ‖₂     (Euclidean distance to cluster i)\nexp = 2 / (m − 1)              (m = fuzziness = 2.0 → exp = 2.0)\nuᵢ = (1/dᵢᵉˣᵖ) / Σⱼ(1/dⱼᵉˣᵖ)  (standard FCM membership)\n\nentropy = −Σᵢ uᵢ·ln(uᵢ) / ln(K)  (normalized Shannon entropy)\nis_ambiguous = (entropy > 0.85)')
label('What It Measures')
body('Parallel fuzzy signal alongside SVM. When memberships are flat (e.g. 0.40/0.35/0.25), the regime is genuinely ambiguous.')
label('Why This Design', RGBColor(29,78,216))
body('Standard FCM formula (Bezdek, 1981) with fuzziness m=2.0. Shannon entropy normalized by ln(K) maps to [0,1]. Threshold 0.85 means membership >85% of max entropy.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('FCM guarantees Σuᵢ=1 and each uᵢ∈(0,1). Normalizing entropy by ln(K) makes threshold portable across cluster counts.')

# ══════════ FORMULA 7 ══════════
heading(7, 'SHAP Normalized Shares', 'Layer 2C — Explainability  |  SIGNAL')
label('Formula')
formula('rawₖ = |SHAP_valueₖ|           (absolute SHAP value per feature)\nshareₖ = rawₖ / Σⱼ rawⱼ          (fractional share, sums to 1.0)\n\nSoft cap: if shareₖ > 0.65 →\n  overflow = shareₖ − 0.65\n  redistribute equally to uncapped features')
label('What It Measures')
body('Fractional importance shares summing to 1.0 for UI charts and downstream cosine similarity.')
label('Why This Design', RGBColor(29,78,216))
body('Absolute values because we care about magnitude (SHAP can be ±). 65% soft cap prevents one feature from dominating visualization. Uses predicted class SHAP, not mean across classes.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Division by total preserves relative ranking. Soft cap with redistribution maintains sum=1.0 invariant.')

# ══════════ FORMULA 8 ══════════
heading(8, 'SHAP Cosine Similarity', 'Layer 2C.5 — Implicit RAG  |  SIGNAL')
label('Formula')
formula('query_vec = [shap_share₁, ..., shap_share₅]\nmean_vec  = cluster_mean_shap[top_cluster]\n\nshap_cosine = clip( (query·mean)/(‖query‖×‖mean‖),  0.0,  1.0 )\n\nDefault = 0.5 when artifact unavailable or zero-norm')
label('What It Measures')
body('Attribution consistency: is the model reasoning the same way it learned for typical cases in this cluster? 1.0=typical, 0.5=neutral, 0.0=atypical.')
label('Why This Design', RGBColor(29,78,216))
body('Cosine is direction-only (ignores magnitude). Clip to [0,1] because negative cosine is meaningless here. Default 0.5 prevents penalization when artifact missing.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Standard method for attribution pattern consistency in XAI literature.')

# ══════════ FORMULA 9 ══════════
heading(9, 'Novelty Score', 'RAG — Retrieval  |  ROUTING')
label('Formula')
formula('query_vec = L2_normalize([x_scaled(5D), shap_shares(5D)]) → 10D\nmax_cosine = max(cosine_sim(query, neighbor) for neighbor in kNN)\nnovelty_score = clip(1.0 − max_cosine, 0.0, 1.0)\n\nis_novel = (novelty_score > 0.40) → RAG suppressed')
label('What It Measures')
body('Distance from any training point. 0.0=identical to training (known). 1.0=fully orthogonal (novel).')
label('Why This Design', RGBColor(29,78,216))
body('Combined 10D vector captures both macro input and SHAP reasoning. 1−max_cosine because closest neighbor is best evidence of precedent. RAG suppressed for novel cases — distant neighbor votes produce fake confidence.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Cosine distance = 1−cosine_similarity is standard. Threshold 0.40 means even closest neighbor shares <60% directional similarity.')

# ══════════ FORMULA 10 ══════════
heading(10, 'RAG Confidence + ARIMA Modifier', 'RAG — Retrieval  |  ROUTING')
label('Formula')
formula('raw_confidence = votes_for_winner / k    (majority vote, k=5)\n\nARIMA modifier:\n  if sarima_trend ≥ 0.65:  confidence = min(1.0, raw + 0.10)\n  elif sarima_trend ≤ 0.35: confidence = raw × 0.70\n  else:                     confidence = raw')
label('What It Measures')
body('Confidence in k-NN majority vote, modulated by sector trend. Improving sectors make precedents more predictive; declining sectors make them less reliable.')
label('Why This Design', RGBColor(29,78,216))
bullet('Amplification (+0.10 additive): improving sector → small bonus')
bullet('Dampening (×0.70 multiplicative): declining sector → proportional penalty')
bullet('Asymmetry intentional: downside uncertainty is more dangerous')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Additive boosts prevent weak-evidence amplification. Multiplicative penalties preserve zero-at-zero. Follows precautionary principle.')

doc.add_page_break()

# ══════════ FORMULA 11 ══════════
heading(11, 'XAI Score', 'Layer 2C — Explainability  |  SIGNAL')
label('Formula')
formula('xai_score = confidence × max(shap_shares)\n\nconfidence  = SVM regime classification probability\nshap_shares = normalized feature importance [5 features]')
label('What It Measures')
body('Explainability index: how focused AND confident the model is. High = confident + one feature dominates.')
label('Why This Design', RGBColor(29,78,216))
body('Multiplication means both must be high. Confident but diffuse → low. Focused but uncertain → low. Captures: "I know what I\'m doing AND I can point to why."')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Product of two [0,1] values stays in [0,1]. Captures conjunction property exactly.')

# ══════════ FORMULA 12 ══════════
heading(12, 'Idea Signal', 'Layer 3 — Idea Scoring  |  SIGNAL')
label('Formula')
formula('base_fit = FIT_TABLE[(regime, business_model, target_segment)]\ndiff_mult = 0.78 + (differentiation_score − 1) × diff_scale[bm]\n\nstage_delta = stage_deltas[bm][stage]\ncomp_delta  = comp_deltas[bm][competitive_intensity]\nreg_delta   = reg_deltas[sector][regulatory_risk] + b2b_relief\nready_delta = (market_readiness − 3) × ready_scale[bm]\n\nif stage ∈ {idea,validation} AND regime ∈ {CONTRACTING,HIGH_FRICTION}:\n    stage_delta ×= 1.35   (hostile regime amplifier)\n\nidea_signal = clip(base_fit × diff_mult + stage_delta + comp_delta\n                   + reg_delta + ready_delta,  0.12,  0.95)')

label('Key Lookup Table (examples)')
add_table(['Regime × Model × Segment', 'Base Fit'], [
    ['GROWTH × saas × b2b', '0.91'],
    ['GROWTH × marketplace × b2c', '0.87'],
    ['EMERGING × saas × b2b', '0.84'],
    ['HIGH_FRICTION × saas × b2b', '0.72'],
    ['CONTRACTING × marketplace × b2c', '0.28'],
])
doc.add_paragraph()
label('BM-Conditional Weights (examples)')
add_table(['BM Type', 'diff_scale', 'Stage: idea', 'Comp: high'], [
    ['marketplace', '0.07', '−0.14', '−0.16'],
    ['saas', '0.13', '−0.07', '−0.08'],
    ['commission', '0.09', '−0.10', '−0.11'],
    ['hardware', '0.12', '−0.15', '−0.10'],
])
doc.add_paragraph()
label('What It Measures')
body('Encodes domain knowledge about which business models thrive in which regimes. PRIMARY driver of output variability across ideas.')
label('Why This Design', RGBColor(29,78,216))
body('Multiplicative differentiation (base × diff_mult) means differentiation matters MORE in favorable regimes and LESS in hostile ones. Additive adjustments shift independently. Hostile regime amplifier (×1.35) captures compounded early-stage survival risk.')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Separates structural fit (multiplicative) from deltas (additive). Mirrors economic reality. Clip to [0.12, 0.95] prevents extremes.')

doc.add_page_break()

# ══════════ FORMULA 13 ══════════
heading(13, 'Intelligent Score (IS)', 'Layer 2D.5 — Routing  |  ROUTING')
label('Formula')
formula('IS = 0.20×S + 0.25×gap_svm + 0.20×μ_fcm + 0.15×arima + 0.20×shap_cosine\n\nCorrelated Trio Discount:\n  if gap_svm>0.80 AND μ_fcm>0.80 AND shap_cosine>0.80:\n    trio = 0.25×gap + 0.20×μ + 0.20×cos\n    IS = IS − trio + trio×0.85   (15% discount)')

label('Signal Components')
add_table(['Signal', 'Weight', 'Source', 'Description'], [
    ['S (regime)', '0.20', 'Lookup', 'GROWTH=0.85, EMERGING=0.65, FRICTION=0.40, CONTRACT=0.20'],
    ['gap_svm', '0.25', 'SVM proba', 'sorted(p)[0]−sorted(p)[1] (probability margin)'],
    ['μ_fcm', '0.20', 'FCM', 'Top-cluster membership ∈[0,1]'],
    ['arima', '0.15', 'SARIMA', 'Normalized sector trajectory ∈[0.15,0.90]'],
    ['shap_cosine', '0.20', 'SHAP+FCM', 'Cosine to cluster mean ∈[0,1]'],
])
doc.add_paragraph()
label('What It Measures')
body('IS is a ROUTING composite — determines which reasoning path (1–7) the ReAct router takes. NOT a decision signal. L4 makes the actual decision.')
label('Why This Design', RGBColor(29,78,216))
bullet('gap_svm highest weight (0.25): captures how decisive classification was')
bullet('arima lowest weight (0.15): only independent signal, but most volatile')
bullet('Trio discount: gap/μ/shap share x_scaled geometry; prevents inflation in easy cases')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Weighted linear combination of [0,1] signals with consistent directionality. Trio discount addresses statistical dependence — without it, IS over-counts the same evidence 3×.')

# ══════════ FORMULA 14 ══════════
heading(14, 'Legacy TAS (Total Assessment Score)', 'Layer 4a — Display Only  |  LEGACY')
label('Formula')
formula('TAS = confidence×0.30 + sarima_trend×0.20 + idea_signal×0.35 + xai_score×0.15\n\nTier: Strong≥0.76 | Moderate≥0.60 | Mixed≥0.44 | Weak<0.44')
p = doc.add_paragraph()
r = p.add_run('⚠️ ZERO decision influence. TAS is preserved for backward compatibility and display only. The real decision is made by the L4 state machine.')
r.bold = True; r.font.color.rgb = RGBColor(153,27,27)
label('Why This Design', RGBColor(29,78,216))
bullet('idea_signal (0.35): highest weight — most idea-specific')
bullet('confidence (0.30): macro classification certainty')
bullet('sarima_trend (0.20): temporal dimension')
bullet('xai_score (0.15): rewards focused explanations')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Weighted average of [0,1] signals stays in [0,1]. Prioritizes idea-specific fit over macro classification.')

# ══════════ FORMULA 15 ══════════
heading(15, 'Decision Quality Assessment', 'Layer 4 — Decision Engine  |  DECISION')
label('Formula')
formula('Input Completeness (ic):\n  agg_conf≥0.75 AND no unknowns → "high"\n  agg_conf≥0.55 AND ≤1 unknown  → "medium"\n  else                          → "low"\n\nSignal Agreement (sa):\n  3 checks: L3 bm available? No insufficient modules? No high conflicts?\n  3/3→"high" | 2/3→"medium" | ≤1/3→"low"\n\nAssumption Density (ad):\n  h_ratio = heuristic / (heuristic + grounded)\n  ≥0.70→"high" | ≥0.50→"medium" | else→"low"\n\nOverall Uncertainty:\n  mech_contribution = mechanism_uncertainty / 0.30\n  bad_tiers = count(ic=low) + count(sa=low) + count(ad=high)\n              + mech_contribution\n  ≥2 OR stale→"high" | ≥1→"moderate" | else→"low"')
label('What It Measures')
body('Replaces numeric confidence with structured, multi-axis quality assessment.')
label('Why This Design', RGBColor(29,78,216))
bullet('Three independent axes prevent hiding dimensional failures')
bullet('Mechanism uncertainty blended continuously — hard-switching loses info')
bullet('mech_uncertainty/0.30 normalizes to [0,1] for proportional contribution')
label('Why It\'s Correct', RGBColor(21,128,61))
body('Continuous blending handles boundary cases correctly. mechanism_uncertainty of 0.20 adds 0.67 — not enough alone for "high," but enough when combined with one bad tier. Preserves ambiguity when ambiguity is real.')

# ══════════ FOOTER ══════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MIDAN Intelligence Engine — Scoring Formulas Reference')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(148,163,184)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Intelligent Systems Project — All 15 formulas in pipeline execution order')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(148,163,184)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MIDAN_Scoring_Formulas.docx')
doc.save(out)
print('DONE: Word document saved successfully!')
